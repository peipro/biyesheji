#!/usr/bin/env python
# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

def clean_text(line):
    """彻底清理markdown格式和所有特殊符号"""
    # 移除所有emoji和特殊符号
    emoji_pattern = re.compile(
        "["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002500-\U00002BEF"
        u"\U0001f300-\U0001f5ff"
        u"\U0001f1e6-\U0001f1ff"
        u"\ufe0f"
        u"\U000020E3"
        "]+",
        flags=re.UNICODE
    )
    line = emoji_pattern.sub(r'', line)

    # 移除变异选择器
    line = re.sub(r'️', '', line)

    # 彻底处理公式：移除所有 $ 和 $$ 符号，保留内容
    line = re.sub(r'\$\$(.*?)\$\$', r'\1', line)
    line = re.sub(r'\$(.*?)\$', r'\1', line)

    # 处理markdown格式
    line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # 粗体
    line = re.sub(r'\*(.*?)\*', r'\1', line)        # 斜体
    line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line) # 链接
    line = re.sub(r'`(.*?)`', r'\1', line)         # 行内代码

    # 移除首尾空白
    return line.strip()

def markdown_to_docx(md_file_path, docx_file_path):
    """改进版Markdown转Word，优化格式"""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)  # 黑色

    # 设置默认段落格式
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_before = Pt(0)

    lines = content.split('\n')
    in_table = False
    table_data = []

    for line in lines:
        line = line.rstrip('\n')
        if not line:
            if in_table and table_data:
                # 创建表格
                if len(table_data) > 0:
                    rows = len(table_data)
                    cols = len(table_data[0]) if rows > 0 else 0
                    if cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'  # 使用完整网格样式，边框黑色
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        # 设置每个单元格
                        for i, row_data in enumerate(table_data):
                            for j, cell_data in enumerate(row_data):
                                cell_data_clean = clean_text(cell_data)
                                para = table.cell(i, j).paragraphs[0]
                                run = para.add_run(cell_data_clean)
                                run.font.name = 'SimHei'
                                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                                if i == 0:  # 表头加粗
                                    run.bold = True
                                # 根据内容长度调整字号
                                if rows > 8:
                                    run.font.size = Pt(9.5)
                                else:
                                    run.font.size = Pt(10)
                        # 添加表格下方间距
                        doc.add_paragraph()
                table_data = []
            in_table = False
            continue

        # 标题处理 - 层级分明
        if line.startswith('# '):
            in_table = False
            text_clean = clean_text(line[2:])
            heading = doc.add_heading(text_clean, level=1)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                run.font.size = Pt(18)
            doc.add_paragraph()
            continue
        elif line.startswith('## '):
            in_table = False
            text_clean = clean_text(line[3:])
            heading = doc.add_heading(text_clean, level=2)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                run.font.size = Pt(15)
            doc.add_paragraph()
            continue
        elif line.startswith('### '):
            in_table = False
            text_clean = clean_text(line[4:])
            heading = doc.add_heading(text_clean, level=3)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                run.font.size = Pt(13)
            continue
        elif line.startswith('#### '):
            in_table = False
            text_clean = clean_text(line[5:])
            heading = doc.add_heading(text_clean, level=4)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                run.font.size = Pt(12)
            continue

        # 分隔线
        elif line.strip() == '---':
            in_table = False
            doc.add_page_break()
            continue

        # 表格行
        elif '|' in line:
            cells = [cell.strip() for cell in line.split('|')]
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            # 跳过markdown表格分隔线（全是-）
            if all(all(c == '-' for c in cell.strip().replace(':', '')) for cell in cells):
                continue
            in_table = True
            table_data.append(cells)
            continue

        # 项目符号列表
        stripped = line.strip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            in_table = False
            text = clean_text(stripped[2:])
            p = doc.add_paragraph('• ' + text)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = 'SimHei'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(11)
            continue

        # 数字编号列表
        elif stripped and stripped[0].isdigit() and '.' in stripped[:2]:
            in_table = False
            dot_pos = stripped.find('.')
            text = clean_text(stripped[dot_pos+1:])
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = 'SimHei'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(11)
            continue

        # 普通段落
        else:
            in_table = False
            text = clean_text(line)
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.name = 'SimHei'
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.size = Pt(11)

    # 处理最后一个表格
    if in_table and table_data:
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        if cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    cell_data_clean = clean_text(cell_data)
                    para = table.cell(i, j).paragraphs[0]
                    run = para.add_run(cell_data_clean)
                    run.font.name = 'SimHei'
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if i == 0:
                        run.bold = True
                    run.font.size = Pt(10)

    doc.save(docx_file_path)
    print(f"转换完成: {docx_file_path}")

if __name__ == '__main__':
    md_file = '实验总结.md'
    docx_file = '实验总结.docx'
    markdown_to_docx(md_file, docx_file)
