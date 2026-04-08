#!/usr/bin/env python
# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor

def clean_text(line):
    """彻底清理markdown格式和emoji"""
    # 移除所有emoji - 扩大范围
    emoji_pattern = re.compile(
        "["
        u"\U0001F600-\U0001F64F"  # emoticons
        u"\U0001F300-\U0001F5FF"  # symbols & pictographs
        u"\U0001F1E0-\U0001F1FF"  # flags
        u"\U00002500-\U00002BEF"  # various symbols
        u"\U0001f300-\U0001f5ff"
        u"\U0001f1e6-\U0001f1ff"
        u"\ufe0f"  # variation selectors
        u"\U000020E3"  # combining enclosing keycap
        "]+",
        flags=re.UNICODE
    )
    line = emoji_pattern.sub(r'', line)

    # 移除编号emoji残留 1️⃣ 变成 1
    line = re.sub(r'️', '', line)

    # 处理公式：$$ ... $$ → 保留内容
    line = re.sub(r'\$\$(.*?)\$\$', r'\1', line)
    # 处理行内公式：$ ... $ → 保留内容
    line = re.sub(r'\$(.*?)\$', r'\1', line)

    # 粗体 **text** → 保留内容
    line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
    # 斜体 *text* → 保留内容
    line = re.sub(r'\*(.*?)\*', r'\1', line)
    # 链接 [text](url) → text
    line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line)
    # 行内代码 `code` → code
    line = re.sub(r'`(.*?)`', r'\1', line)
    # 移除多余空格
    line = line.strip()
    return line

def markdown_to_docx(md_file_path, docx_file_path):
    """彻底改进的Markdown转Word"""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    lines = content.split('\n')
    in_table = False
    table_data = []

    for line in lines:
        line = line.rstrip('\n')
        if not line:
            if in_table and table_data:
                # 创建表格
                if len(table_data) > 0:
                    rows = len(table_data)
                    cols = len(table_data[0]) if rows > 0 else 0
                    if cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Light Grid Accent 1'
                        # 设置表格字体
                        for i, row_data in enumerate(table_data):
                            for j, cell_data in enumerate(row_data):
                                cell_data_clean = clean_text(cell_data)
                                para = table.cell(i, j).paragraphs[0]
                                run = para.add_run(cell_data_clean)
                                run.font.name = 'SimHei'
                                run.font.size = Pt(10)
                        doc.add_paragraph()
                table_data = []
            in_table = False
            continue

        # 标题
        if line.startswith('# '):
            in_table = False
            text_clean = clean_text(line[2:])
            heading = doc.add_heading(text_clean, level=1)
            for run in heading.runs:
                run.font.name = 'SimHei'
            continue
        elif line.startswith('## '):
            in_table = False
            text_clean = clean_text(line[3:])
            heading = doc.add_heading(text_clean, level=2)
            for run in heading.runs:
                run.font.name = 'SimHei'
            doc.add_paragraph()
            continue
        elif line.startswith('### '):
            in_table = False
            text_clean = clean_text(line[4:])
            heading = doc.add_heading(text_clean, level=3)
            for run in heading.runs:
                run.font.name = 'SimHei'
            continue
        elif line.startswith('#### '):
            in_table = False
            text_clean = clean_text(line[5:])
            heading = doc.add_heading(text_clean, level=4)
            for run in heading.runs:
                run.font.name = 'SimHei'
            continue

        # 分隔线
        elif line.strip() == '---':
            in_table = False
            doc.add_paragraph()
            continue

        # 表格行
        elif '|' in line:
            cells = [cell.strip() for cell in line.split('|')]
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            # 跳过表格分隔线（全是-）
            if all(all(c == '-' for c in cell.strip().replace(':', '')) for cell in cells):
                continue
            in_table = True
            table_data.append(cells)
            continue

        # 列表项
        stripped = line.strip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            in_table = False
            text = clean_text(stripped[2:])
            p = doc.add_paragraph('• ' + text)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = 'SimHei'
                run.font.size = Pt(11)
            continue

        elif stripped and stripped[0].isdigit() and '.' in stripped[:2]:
            in_table = False
            dot_pos = stripped.find('.')
            text = clean_text(stripped[dot_pos+1:])
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = 'SimHei'
                run.font.size = Pt(11)
            continue

        # 普通段落
        else:
            in_table = False
            text = clean_text(line)
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.name = 'SimHei'
                    run.font.size = Pt(11)

    # 处理最后一个表格
    if in_table and table_data:
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        if cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    cell_data_clean = clean_text(cell_data)
                    para = table.cell(i, j).paragraphs[0]
                    run = para.add_run(cell_data_clean)
                    run.font.name = 'SimHei'
                    run.font.size = Pt(10)

    doc.save(docx_file_path)
    print(f"转换完成: {docx_file_path}")

if __name__ == '__main__':
    md_file = '实验总结.md'
    docx_file = '实验总结.docx'
    markdown_to_docx(md_file, docx_file)
