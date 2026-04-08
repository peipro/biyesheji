#!/usr/bin/env python
# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def markdown_to_docx(md_file_path, docx_file_path):
    """简单的Markdown转Word"""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimHei'
    font.size = Pt(12)

    lines = content.split('\n')

    for line in lines:
        line = line.rstrip()
        if not line:
            continue

        # 标题
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            continue
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            continue
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            continue

        # 分隔线
        elif line.startswith('---'):
            doc.add_page_break()
            continue

        # 表格处理 - 简单处理，不保存格式，直接添加文本
        elif '|' in line and ('---' not in line):
            # 去掉首尾的|，分割
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_text = '\t'.join(cells)
            p = doc.add_paragraph(table_text)
            continue

        # 跳过表格分隔线
        elif '|' in line and '---' in line:
            continue

        # 列表项
        elif line.startswith('- ') or line.startswith('1. '):
            p = doc.add_paragraph(line, style='List Bullet')
            continue

        # 普通段落
        else:
            # 移除markdown格式符号（简单处理）
            # 粗体 **text** → text
            line_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # 斜体 *text* → text
            line_clean = re.sub(r'\*(.*?)\*', r'\1', line_clean)
            # 链接 [text](url) → text
            line_clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line_clean)
            # 代码 `code` → code
            line_clean = re.sub(r'`(.*?)`', r'\1', line_clean)
            # emoji
            line_clean = re.sub(r'[🟢🔵⚫⚪⚠️✅⏳️1️⃣2️⃣3️⃣4️⃣5️⃣6️⃣7️⃣8️⃣9️⃣]', '', line_clean)

            p = doc.add_paragraph(line_clean)

    doc.save(docx_file_path)
    print(f"转换完成: {docx_file_path}")

if __name__ == '__main__':
    md_file = '实验总结.md'
    docx_file = '实验总结.docx'
    markdown_to_docx(md_file, docx_file)
